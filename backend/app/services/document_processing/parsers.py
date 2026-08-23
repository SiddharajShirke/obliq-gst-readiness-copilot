"""Deterministic PDF, image, CSV, Excel and JSON parsers."""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

import pandas as pd

from app.schemas.documents import NormalizedGSTRecord

COLUMN_ALIASES = {
    "invoice_number": {
        "invoice no",
        "invoice number",
        "invoice_no",
        "inv no",
        "bill no",
        "document number",
        "document no",
        "document no.",
        "bill of entry no",
        "bill of entry no.",
    },
    "invoice_date": {"invoice date", "invoice_date", "date", "bill date", "document date"},
    "supplier_name": {"supplier", "supplier name", "vendor", "vendor name", "party name"},
    "supplier_gstin": {"supplier gstin", "vendor gstin", "gstin", "supplier_gst_no"},
    "customer_name": {
        "customer",
        "customer name",
        "buyer",
        "buyer name",
        "customer / recipient",
    },
    "customer_gstin": {"customer gstin", "buyer gstin", "recipient gstin"},
    "taxable_value": {
        "taxable amount",
        "taxable value",
        "taxable_value",
        "assessable value",
        "taxable",
    },
    "cgst": {"cgst", "cgst amount", "cgst_amt", "cgst charged"},
    "sgst": {
        "sgst",
        "sgst amount",
        "sgst_amt",
        "utgst",
        "sgst/utgst",
        "sgst charged",
    },
    "igst": {"igst", "igst amount", "igst_amt", "igst charged"},
    "cess": {"cess", "cess amount"},
    "invoice_total": {"total", "invoice total", "invoice value", "grand total", "gross amount"},
    "gst_rate": {"gst rate", "tax rate", "rate %", "gst %"},
    "itc_status": {"itc status", "itc availability", "itc eligibility", "itc intent", "itc"},
    "rcm_flag": {"rcm", "rcm flag", "reverse charge", "reverse charge flag"},
    "rcm_tax_liability": {"rcm tax liability"},
    "transaction_type": {
        "transaction type",
        "supply type",
        "purchase/expense type",
        "note type",
        "doc type",
        "document type",
        "purchase category",
        "supply category",
    },
    "original_document_reference": {
        "original invoice reference",
        "original document reference",
        "original invoice",
    },
    "place_of_supply": {"place of supply", "pos"},
    "hsn_sac": {"hsn/sac", "hsn", "sac"},
}


@dataclass(slots=True)
class ParsedTable:
    rows: list[dict[str, Any]]
    summary: dict[str, Any]
    column_mapping: dict[str, str]


@dataclass(slots=True)
class ParsedNormalizedTable:
    records: list[NormalizedGSTRecord]
    summary: dict[str, Any]
    column_mapping: dict[str, str]


def _normalize_header(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value).strip().lower().replace("_", " "))


def map_columns(columns: list[Any]) -> dict[str, str]:
    normalized = {_normalize_header(column): str(column) for column in columns}
    mapping: dict[str, str] = {}
    for canonical, aliases in COLUMN_ALIASES.items():
        for alias in aliases | {canonical.replace("_", " ")}:
            if alias in normalized:
                mapping[canonical] = normalized[alias]
                break
    return mapping


def _to_float(value: Any) -> float:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return 0.0
    text = str(value).replace(",", "").replace("₹", "").strip()
    try:
        return float(text)
    except ValueError:
        return 0.0


def _to_decimal_optional(value: Any) -> Decimal | None:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None
    text = str(value).replace(",", "").replace("₹", "").strip()
    if not text:
        return None
    try:
        return Decimal(text).quantize(Decimal("0.01"))
    except (InvalidOperation, ValueError):
        return None


def _to_bool_optional(value: Any) -> bool | None:
    text = str(value or "").strip().lower()
    if text in {"yes", "y", "true", "1", "applicable"}:
        return True
    if text in {"no", "n", "false", "0", "not applicable"}:
        return False
    return None


def _mapped_value(source: Any, mapping: dict[str, str], field: str) -> Any:
    column = mapping.get(field)
    return source.get(column) if column else None


def _to_iso_date(value: Any) -> str | None:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None
    if isinstance(value, (date, datetime, pd.Timestamp)):
        return value.date().isoformat() if isinstance(value, datetime) else value.isoformat()
    text = str(value).strip()
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", text):
        try:
            return date.fromisoformat(text).isoformat()
        except ValueError:
            return None
    parsed = pd.to_datetime(text, errors="coerce", dayfirst=True)
    return None if pd.isna(parsed) else parsed.date().isoformat()


def parse_tabular_document(content: bytes, extension: str, *, category: str) -> ParsedTable:
    extension = extension.lower()
    if extension == ".csv":
        frame = pd.read_csv(io.BytesIO(content))
    elif extension in {".xlsx", ".xls"}:
        frame = pd.read_excel(io.BytesIO(content))
    elif extension == ".json":
        payload = json.loads(content.decode("utf-8"))
        if isinstance(payload, dict):
            payload = payload.get("records") or payload.get("invoices") or payload.get("b2b") or []
        frame = pd.DataFrame(payload)
    else:
        raise ValueError(f"Unsupported tabular extension: {extension}")

    mapping = map_columns(list(frame.columns))
    if "invoice_number" not in mapping:
        raise ValueError("Could not identify an invoice-number column")

    rows: list[dict[str, Any]] = []
    for _, source in frame.fillna("").iterrows():
        record: dict[str, Any] = {
            "invoice_category": category,
            "invoice_number": str(source.get(mapping["invoice_number"], "")).strip(),
            "invoice_date": _to_iso_date(source.get(mapping.get("invoice_date", ""))),
            "supplier_name": str(source.get(mapping.get("supplier_name", ""), "")).strip() or None,
            "supplier_gstin": str(source.get(mapping.get("supplier_gstin", ""), "")).strip().upper()
            or None,
            "customer_name": str(source.get(mapping.get("customer_name", ""), "")).strip() or None,
            "customer_gstin": str(source.get(mapping.get("customer_gstin", ""), "")).strip().upper()
            or None,
        }
        for field in ("taxable_value", "cgst", "sgst", "igst", "cess", "invoice_total"):
            record[field] = _to_float(source.get(mapping.get(field, ""), 0))
        rows.append(record)

    summary = {
        "invoice_count": len(rows),
        "taxable_value": round(sum(row["taxable_value"] for row in rows), 2),
        "cgst": round(sum(row["cgst"] for row in rows), 2),
        "sgst": round(sum(row["sgst"] for row in rows), 2),
        "igst": round(sum(row["igst"] for row in rows), 2),
        "cess": round(sum(row["cess"] for row in rows), 2),
        "invoice_total": round(sum(row["invoice_total"] for row in rows), 2),
        "missing_gstin_count": sum(not row.get("supplier_gstin") for row in rows),
    }
    return ParsedTable(rows=rows, summary=summary, column_mapping=mapping)


def parse_normalized_table(
    content: bytes,
    extension: str,
    *,
    document_type: str,
    source_document_id: str,
    tax_period: str | None,
) -> ParsedNormalizedTable:
    extension = extension.lower()
    if extension == ".csv":
        frame = pd.read_csv(io.BytesIO(content))
    elif extension in {".xlsx", ".xls"}:
        frame = pd.read_excel(io.BytesIO(content))
    elif extension == ".json":
        payload = json.loads(content.decode("utf-8-sig"))
        if isinstance(payload, dict):
            payload = payload.get("records") or payload.get("invoices") or payload.get("b2b") or []
        frame = pd.DataFrame(payload)
    else:
        raise ValueError(f"Unsupported tabular extension: {extension}")
    records, mapping = _normalized_records_from_frame(
        frame,
        document_type=document_type,
        source_document_id=source_document_id,
        tax_period=tax_period,
    )
    return _parsed_normalized_table(records, mapping)


def _normalized_records_from_frame(
    frame: pd.DataFrame,
    *,
    document_type: str,
    source_document_id: str,
    tax_period: str | None,
    source_page: int | None = None,
) -> tuple[list[NormalizedGSTRecord], dict[str, str]]:
    mapping = map_columns(list(frame.columns))
    if "invoice_number" not in mapping:
        raise ValueError("Could not identify a document-number column")

    records: list[NormalizedGSTRecord] = []
    for index, source in frame.iterrows():

        def value(field: str, source_row: Any = source) -> Any:
            return _mapped_value(source_row, mapping, field)

        taxes = {
            "igst": _to_decimal_optional(value("igst")),
            "cgst": _to_decimal_optional(value("cgst")),
            "sgst_utgst": _to_decimal_optional(value("sgst")),
            "cess": _to_decimal_optional(value("cess")),
        }
        present_taxes = [amount for amount in taxes.values() if amount is not None]
        total_tax = sum(present_taxes, Decimal("0")) if present_taxes else None
        transaction_type = str(value("transaction_type") or "").strip().lower() or None
        rcm_flag = _to_bool_optional(value("rcm_flag"))
        rcm_liability = _to_decimal_optional(value("rcm_tax_liability"))
        if rcm_flag is None and (
            (rcm_liability is not None and rcm_liability != 0)
            or "rcm" in str(transaction_type or "")
        ):
            rcm_flag = True
        record = NormalizedGSTRecord(
            tax_period=tax_period,
            document_type=document_type,
            document_number=str(value("invoice_number") or "").strip() or None,
            document_date=_to_iso_date(value("invoice_date")),
            supplier_name=str(value("supplier_name") or "").strip() or None,
            supplier_gstin=str(value("supplier_gstin") or "").strip().upper() or None,
            customer_name=str(value("customer_name") or "").strip() or None,
            customer_gstin=str(value("customer_gstin") or "").strip().upper() or None,
            place_of_supply=str(value("place_of_supply") or "").strip() or None,
            hsn_sac=str(value("hsn_sac") or "").strip() or None,
            taxable_value=_to_decimal_optional(value("taxable_value")),
            gst_rate=_to_decimal_optional(value("gst_rate")),
            total_document_value=_to_decimal_optional(value("invoice_total")),
            transaction_type=transaction_type,
            itc_status=str(value("itc_status") or "").strip().lower() or None,
            rcm_flag=rcm_flag,
            original_document_reference=(
                str(value("original_document_reference") or "").strip() or None
            ),
            source_document_id=source_document_id,
            source_page=source_page,
            source_row=int(index) + 2,
            total_tax=total_tax,
            **taxes,
        )
        records.append(record)

    return records, mapping


def _parsed_normalized_table(
    records: list[NormalizedGSTRecord], mapping: dict[str, str]
) -> ParsedNormalizedTable:

    def money(field: str) -> Decimal:
        return sum(
            (getattr(record, field) or Decimal("0") for record in records),
            Decimal("0"),
        )

    return ParsedNormalizedTable(
        records=records,
        summary={
            "record_count": len(records),
            "taxable_value": str(money("taxable_value")),
            "total_tax": str(money("total_tax")),
            "total_document_value": str(money("total_document_value")),
            "rcm_count": sum(record.rcm_flag is True for record in records),
            "itc_not_available_count": sum(
                record.itc_status in {"not available", "ineligible", "blocked"}
                for record in records
            ),
        },
        column_mapping=mapping,
    )


def _clean_pdf_table_cell(value: Any) -> str:
    """Remove isolated watermark letters while preserving real table text."""

    lines = [line.strip() for line in str(value or "").splitlines()]
    meaningful = [line for line in lines if not re.fullmatch(r"[A-Z]", line)]
    return " ".join(line for line in meaningful if line).strip()


def parse_normalized_pdf_tables(
    content: bytes,
    *,
    document_type: str,
    source_document_id: str,
    tax_period: str | None,
) -> ParsedNormalizedTable | None:
    """Extract structured register/GSTR-2B tables before considering an LLM."""

    import pymupdf

    records: list[NormalizedGSTRecord] = []
    combined_mapping: dict[str, str] = {}
    with pymupdf.open(stream=content, filetype="pdf") as document:
        for page_number, page in enumerate(document, start=1):
            for table in page.find_tables().tables:
                extracted = table.extract()
                if len(extracted) < 2:
                    continue
                headers = [_clean_pdf_table_cell(value) for value in extracted[0]]
                frame = pd.DataFrame(
                    [
                        [_clean_pdf_table_cell(value) for value in row]
                        for row in extracted[1:]
                    ],
                    columns=headers,
                )
                mapping = map_columns(headers)
                if "invoice_number" not in mapping or not {
                    "taxable_value",
                    "invoice_total",
                }.intersection(mapping):
                    continue
                table_records, table_mapping = _normalized_records_from_frame(
                    frame,
                    document_type=document_type,
                    source_document_id=source_document_id,
                    tax_period=tax_period,
                    source_page=page_number,
                )
                records.extend(
                    record for record in table_records if record.document_number is not None
                )
                combined_mapping.update(table_mapping)
    return _parsed_normalized_table(records, combined_mapping) if records else None


def read_pdf_text(content: bytes) -> str:
    import pymupdf

    with pymupdf.open(stream=content, filetype="pdf") as document:
        return "\n".join(page.get_text("text") for page in document).strip()


def read_image_text(content: bytes, *, tesseract_cmd: str = "") -> str:
    import pytesseract
    from PIL import Image

    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    image = Image.open(io.BytesIO(content))
    return pytesseract.image_to_string(image).strip()


def read_scanned_pdf_text(content: bytes, *, tesseract_cmd: str = "") -> str:
    """OCR PDF pages with the deployment-packaged Tesseract binary."""
    import pymupdf

    pages: list[str] = []
    with pymupdf.open(stream=content, filetype="pdf") as document:
        for page in document:
            pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
            pages.append(read_image_text(pixmap.tobytes("png"), tesseract_cmd=tesseract_cmd))
    return "\n".join(value for value in pages if value).strip()


def read_docx_text(content: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(content))
    return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()


def _label(text: str, labels: list[str]) -> str | None:
    escaped = "|".join(re.escape(label) for label in labels)
    match = re.search(rf"(?:{escaped})\s*[:\-]\s*([^\n\r]+)", text, flags=re.IGNORECASE)
    return match.group(1).strip() if match else None


def _amount(text: str, labels: list[str]) -> float:
    raw = _label(text, labels)
    if not raw:
        return 0.0
    number = re.search(r"[-+]?\d[\d,]*(?:\.\d+)?", raw.replace("₹", ""))
    return float(number.group(0).replace(",", "")) if number else 0.0


def _invoice_date(text: str) -> str | None:
    raw = _label(text, ["Invoice Date", "Bill Date", "Date"])
    if not raw:
        return None
    parsed = pd.to_datetime(raw, errors="coerce", dayfirst=True)
    return None if pd.isna(parsed) else parsed.date().isoformat()


def extract_invoice_from_text(text: str, document_type: str) -> dict[str, Any]:
    result = {
        "document_type": document_type,
        "supplier_name": _label(text, ["Supplier", "Vendor", "Seller"]),
        "supplier_gstin": (
            _label(text, ["Supplier GSTIN", "Vendor GSTIN", "Seller GSTIN"]) or ""
        ).upper()
        or None,
        "customer_name": _label(text, ["Customer", "Buyer", "Recipient"]),
        "customer_gstin": (
            _label(text, ["Customer GSTIN", "Buyer GSTIN", "Recipient GSTIN"]) or ""
        ).upper()
        or None,
        "invoice_number": _label(text, ["Invoice Number", "Invoice No", "Bill No"]),
        "invoice_date": _invoice_date(text),
        "place_of_supply": _label(text, ["Place of Supply"]),
        "taxable_value": _amount(text, ["Taxable Value", "Taxable Amount"]),
        "cgst": _amount(text, ["CGST"]),
        "sgst": _amount(text, ["SGST", "UTGST"]),
        "igst": _amount(text, ["IGST"]),
        "cess": _amount(text, ["Cess"]),
        "invoice_total": _amount(text, ["Invoice Total", "Grand Total", "Total"]),
        "hsn_sac": _label(text, ["HSN", "SAC", "HSN/SAC"]),
        "line_items": [],
        "field_confidences": {},
        "overall_confidence": 0.9,
        "warnings": [],
    }
    present = sum(
        result.get(key) not in (None, "", 0.0)
        for key in (
            "supplier_name",
            "supplier_gstin",
            "invoice_number",
            "invoice_date",
            "taxable_value",
            "invoice_total",
        )
    )
    result["overall_confidence"] = round(0.45 + (present / 6) * 0.5, 2)
    return result


def parse_document_content(
    filename: str, document_type: str, mime_type: str, content: bytes, *, tesseract_cmd: str = ""
) -> tuple[str, dict[str, Any]]:
    extension = Path(filename).suffix.lower()
    if document_type in {"sales_register", "purchase_register", "gstr2b"}:
        category = {"sales_register": "sales", "purchase_register": "purchase", "gstr2b": "gstr2b"}[
            document_type
        ]
        table = parse_tabular_document(content, extension, category=category)
        return "", {
            "summary": table.summary,
            "rows": table.rows,
            "column_mapping": table.column_mapping,
        }
    if extension == ".pdf":
        text = read_pdf_text(content)
    elif extension in {".png", ".jpg", ".jpeg"}:
        text = read_image_text(content, tesseract_cmd=tesseract_cmd)
    else:
        text = content.decode("utf-8", errors="ignore")
    return text, extract_invoice_from_text(text, document_type)
